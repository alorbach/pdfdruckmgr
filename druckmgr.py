#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Print Manager - A tool for double-sided image printing
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk, ImageChops
import os
import sys
from pathlib import Path
import subprocess
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import io

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False
    print("Warning: tkinterdnd2 not available. Install with: pip install tkinterdnd2")
    print("Drag & Drop wird nicht funktionieren.")


class DruckManager:
    def __init__(self, root):
        self.root = root
        self.root.title("Print Manager - Double-Sided Image Printing")
        self.root.geometry("1200x800")
        
        # Variablen
        self.images = []  # Liste von (vorderseite, rueckseite) Tupeln
        self.image_mirrors = {}  # {(pair_index, 'front'/'back'): ('h'/'v'/'both'/'none')}
        self.current_pair_index = 0
        self.debug_mode = tk.BooleanVar(value=False)
        self.mirror_back = tk.BooleanVar(value=True)
        self.margin = tk.DoubleVar(value=1.0)  # in cm
        self.scale_to_width = tk.BooleanVar(value=True)
        self.auto_open_export = tk.BooleanVar(value=True)  # Exportierte Dateien automatisch oeffnen
        self.pdf_landscape = tk.BooleanVar(value=False)  # PDF im Querformat (default aus)
        self.auto_trim = tk.BooleanVar(value=True)  # Weissen Rand automatisch entfernen
        self.target_width = 29.7  # cm (A4 Breite)
        
        # Drag & Drop Variablen
        self.drag_start_index = None
        self.drag_start_y = None
        self.dragging = False
        self.drag_threshold = 5  # Pixel, um zwischen Klick und Drag zu unterscheiden
        
        # Bild-Tausch Variablen
        self.image_drag_source = None  # (pair_index, 'front'/'back')
        self.image_dragging = False
        
        # Debug Ausgabe
        self.debug_text = None
        
        self.setup_ui()
        self.log_debug("Anwendung gestartet")
        
    def setup_ui(self):
        # Hauptcontainer
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        # Linke Seite - Drag & Drop Bereich
        left_frame = ttk.LabelFrame(main_frame, text="Add images", padding="10")
        left_frame.grid(row=0, column=0, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        # Drag & Drop Bereich
        self.drop_area = tk.Text(left_frame, width=30, height=15, relief=tk.SUNKEN, 
                                 bg="lightgray", wrap=tk.WORD)
        self.drop_area.pack(fill=tk.BOTH, expand=True)
        self.drop_area.insert("1.0", "Drag images here\n\nOr click 'Select images'")
        self.drop_area.config(state=tk.DISABLED)
        
        if DND_AVAILABLE:
            self.drop_area.drop_target_register(DND_FILES)
            self.drop_area.dnd_bind('<<Drop>>', self.on_drop)
        
        # Button zum manuellen Auswaehlen
        ttk.Button(left_frame, text="Select images", 
                  command=self.select_images).pack(pady=5)
        
        ttk.Button(left_frame, text="Clear all", 
                  command=self.clear_all).pack(pady=5)
        
        # Mitte - Vorschau
        middle_frame = ttk.LabelFrame(main_frame, text="Preview", padding="10")
        middle_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5)
        
        # Vorderseite Vorschau
        ttk.Label(middle_frame, text="Front:").pack()
        front_frame = tk.Frame(middle_frame, bg="white", relief=tk.SUNKEN, borderwidth=2)
        front_frame.pack(pady=5, fill=tk.BOTH, expand=True)
        self.front_preview = tk.Label(front_frame, bg="white")
        self.front_preview.pack(expand=True, fill=tk.BOTH, padx=2, pady=2)
        
        # Rueckseite Vorschau
        ttk.Label(middle_frame, text="Back:").pack()
        back_frame = tk.Frame(middle_frame, bg="white", relief=tk.SUNKEN, borderwidth=2)
        back_frame.pack(pady=5, fill=tk.BOTH, expand=True)
        self.back_preview = tk.Label(back_frame, bg="white")
        self.back_preview.pack(expand=True, fill=tk.BOTH, padx=2, pady=2)
        
        # Navigation
        nav_frame = ttk.Frame(middle_frame)
        nav_frame.pack(pady=10)
        ttk.Button(nav_frame, text="< Previous", 
                  command=self.prev_pair).pack(side=tk.LEFT, padx=5)
        self.pair_label = ttk.Label(nav_frame, text="Pair 0 of 0")
        self.pair_label.pack(side=tk.LEFT, padx=5)
        ttk.Button(nav_frame, text="Next >", 
                  command=self.next_pair).pack(side=tk.LEFT, padx=5)
        
        # Tausch Button
        ttk.Button(middle_frame, text="↔ Swap front/back", 
                  command=lambda: self.swap_pair_images(self.current_pair_index)).pack(pady=5)
        
        # Rechte Seite - Kachelansicht
        right_frame = ttk.LabelFrame(main_frame, text="Tiles", padding="10")
        right_frame.grid(row=0, column=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        # Scrollbar fuer Kachelansicht
        canvas_frame = ttk.Frame(right_frame)
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        self.tile_canvas = tk.Canvas(canvas_frame, bg="white")
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.tile_canvas.yview)
        self.tile_scrollable = ttk.Frame(self.tile_canvas)
        
        self.tile_scrollable.bind(
            "<Configure>",
            lambda e: self.tile_canvas.configure(scrollregion=self.tile_canvas.bbox("all"))
        )
        
        self.tile_canvas.create_window((0, 0), window=self.tile_scrollable, anchor="nw")
        self.tile_canvas.configure(yscrollcommand=scrollbar.set)
        
        self.tile_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Einstellungen
        settings_frame = ttk.LabelFrame(main_frame, text="Settings", padding="10")
        settings_frame.grid(row=1, column=1, columnspan=2, sticky=(tk.W, tk.E), padx=5, pady=(5, 0))
        
        # Margen
        margin_frame = ttk.Frame(settings_frame)
        margin_frame.pack(fill=tk.X, pady=5)
        ttk.Label(margin_frame, text="Margins (cm):").pack(side=tk.LEFT, padx=5)
        margin_spin = ttk.Spinbox(margin_frame, from_=0.0, to=5.0, increment=0.1, 
                                  textvariable=self.margin, width=10,
                                  command=lambda: self.update_previews())
        margin_spin.pack(side=tk.LEFT, padx=5)
        self.margin.trace_add("write", lambda *args: self.update_previews())
        
        # Spiegeln
        ttk.Checkbutton(settings_frame, text="Mirror back side automatically", 
                       variable=self.mirror_back,
                       command=self.update_previews).pack(anchor=tk.W, pady=5)
        
        # Skalierung
        ttk.Checkbutton(settings_frame, text="Scale to A4 width (29.7 cm)", 
                       variable=self.scale_to_width,
                       command=self.update_previews).pack(anchor=tk.W, pady=5)

        # Rand entfernen
        ttk.Checkbutton(settings_frame, text="Auto trim white borders", 
                       variable=self.auto_trim).pack(anchor=tk.W, pady=5)

        # PDF Querformat
        ttk.Checkbutton(settings_frame, text="PDF landscape (A4)", 
                       variable=self.pdf_landscape).pack(anchor=tk.W, pady=5)
        
        # Debug
        ttk.Checkbutton(settings_frame, text="Enable debug output", 
                       variable=self.debug_mode, 
                       command=self.toggle_debug).pack(anchor=tk.W, pady=5)
        
        # Auto-Open Export
        ttk.Checkbutton(settings_frame, text="Auto open exported files", 
                       variable=self.auto_open_export).pack(anchor=tk.W, pady=5)
        
        # Aktionen
        action_frame = ttk.LabelFrame(main_frame, text="Actions", padding="10")
        action_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 0))
        
        ttk.Button(action_frame, text="Print", 
                  command=self.print_images).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Save as PDF", 
                  command=self.export_pdf).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Save as Word", 
                  command=self.export_word).pack(side=tk.LEFT, padx=5)
        
        # Debug Ausgabe (versteckt standardmaessig)
        self.debug_frame = ttk.LabelFrame(main_frame, text="Debug output", padding="10")
        self.debug_text = scrolledtext.ScrolledText(self.debug_frame, height=8, width=100)
        self.debug_text.pack(fill=tk.BOTH, expand=True)
        self.debug_frame.grid_remove()
        
        # Grid Gewichtungen
        main_frame.columnconfigure(1, weight=1)
        main_frame.columnconfigure(2, weight=1)
        main_frame.rowconfigure(0, weight=1)
        
    def log_debug(self, message):
        """Debug Nachricht ausgeben"""
        if self.debug_mode.get() and self.debug_text:
            self.debug_text.insert(tk.END, f"[DEBUG] {message}\n")
            self.debug_text.see(tk.END)
        if self.debug_mode.get():
            print(f"[DEBUG] {message}")
    
    def toggle_debug(self):
        """Debug Frame ein/ausblenden"""
        if self.debug_mode.get():
            self.debug_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 0))
        else:
            self.debug_frame.grid_remove()
    
    def on_drop(self, event):
        """Drag & Drop Handler"""
        files = self.root.tk.splitlist(event.data)
        image_files = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
        self.log_debug(f"Files received via drag & drop: {len(image_files)}")
        self.process_images(image_files)
    
    def select_images(self):
        """Select images manually"""
        files = filedialog.askopenfilenames(
            title="Select images",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All files", "*.*")]
        )
        if files:
            self.log_debug(f"Files selected: {len(files)}")
            self.process_images(list(files))
    
    def process_images(self, image_files):
        """Process images and build pairs"""
        if not image_files:
            return
        
        # Wenn ungerade Anzahl, letztes Bild als einzelnes Paar (Rueckseite leer)
        for i in range(0, len(image_files), 2):
            front = image_files[i]
            back = image_files[i + 1] if i + 1 < len(image_files) else None
            self.images.append((front, back))
            self.log_debug(f"Pair added: front={front}, back={back}")
        
        self.update_previews()
        self.update_tile_view()
        self.log_debug(f"Total pairs: {len(self.images)}")
    
    def clear_all(self):
        """Clear all images"""
        self.images = []
        self.image_mirrors = {}
        self.current_pair_index = 0
        self.update_previews()
        self.update_tile_view()
        self.log_debug("All images cleared")
    
    def update_previews(self):
        """Update preview"""
        if not self.images:
            self.front_preview.config(image='')
            self.back_preview.config(image='')
            if self.pair_label:
                self.pair_label.config(text="Pair 0 of 0")
            return
        
        if 0 <= self.current_pair_index < len(self.images):
            front_path, back_path = self.images[self.current_pair_index]
            
            # Vorderseite
            if front_path:
                self.show_preview(front_path, self.front_preview, 
                                pair_index=self.current_pair_index, side='front')
                # Rechtsklick Menue
                self.front_preview.bind("<Button-3>", lambda e: self.show_image_menu(e, self.current_pair_index, 'front'))
            else:
                self.front_preview.config(image='')
                self.front_preview.unbind("<Button-3>")
            
            # Rueckseite
            if back_path:
                self.show_preview(back_path, self.back_preview,
                                pair_index=self.current_pair_index, side='back')
                # Rechtsklick Menue
                self.back_preview.bind("<Button-3>", lambda e: self.show_image_menu(e, self.current_pair_index, 'back'))
            else:
                self.back_preview.config(image='')
                self.back_preview.unbind("<Button-3>")
            
            if self.pair_label:
                self.pair_label.config(text=f"Pair {self.current_pair_index + 1} of {len(self.images)}")
    
    def show_preview(self, image_path, label_widget, max_size=(400, 300), pair_index=None, side=None):
        """Show image in label"""
        try:
            img = Image.open(image_path)
            
            # Spiegelung anwenden wenn vorhanden
            if pair_index is not None and side is not None:
                mirror = self.image_mirrors.get((pair_index, side), 'none')
                img = self.apply_mirror(img, mirror)
            
            # Thumbnail erstellen mit Seitenverhaeltnis
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            label_widget.config(image=photo)
            label_widget.image = photo  # Referenz behalten
        except Exception as e:
            self.log_debug(f"Failed to load {image_path}: {e}")
            messagebox.showerror("Error", f"Could not load image: {e}")
    
    def apply_mirror(self, img, mirror_type):
        """Spiegelung auf Bild anwenden"""
        if mirror_type == 'h' or mirror_type == 'horizontal':
            return img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
        elif mirror_type == 'v' or mirror_type == 'vertical':
            return img.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
        elif mirror_type == 'both':
            img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
            return img.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
        return img

    def trim_image(self, img):
        """Trim white/transparent borders"""
        try:
            if img.mode in ("RGBA", "LA"):
                alpha = img.split()[-1]
                bbox = alpha.getbbox()
                if bbox:
                    return img.crop(bbox)
            # Fallback: weissen Hintergrund entfernen
            rgb = img.convert("RGB")
            bg = Image.new("RGB", rgb.size, (255, 255, 255))
            diff = ImageChops.difference(rgb, bg)
            bbox = diff.getbbox()
            if bbox:
                return img.crop(bbox)
        except Exception as e:
            self.log_debug(f"Trim failed: {e}")
        return img

    def load_base_image(self, image_path, pair_index=None, side=None, mirror=False, trim=False):
        """Bild laden, spiegeln und optional zuschneiden (ohne Skalierung)"""
        img = Image.open(image_path)
        if pair_index is not None and side is not None:
            mirror_type = self.image_mirrors.get((pair_index, side), 'none')
            img = self.apply_mirror(img, mirror_type)
        elif mirror:
            img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
        if trim:
            img = self.trim_image(img)
        return img

    def compute_target_size_cm(self, img, available_width_cm, available_height_cm):
        """Zielgroesse in cm anhand des Seitenverhaeltnisses berechnen"""
        aspect_ratio = img.height / img.width
        width_cm = available_width_cm
        height_cm = width_cm * aspect_ratio
        if height_cm > available_height_cm:
            height_cm = available_height_cm
            width_cm = height_cm / aspect_ratio
        return width_cm, height_cm

    def prepare_export_image(self, img, target_width_cm=None, target_height_cm=None):
        """Scale image for export (PDF/Word) to target size"""
        if target_width_cm is None or target_height_cm is None:
            return img
        target_dpi = 300
        target_width_px = int(target_width_cm / 2.54 * target_dpi)
        target_height_px = int(target_height_cm / 2.54 * target_dpi)
        return img.resize((target_width_px, target_height_px), Image.Resampling.LANCZOS)
    
    def update_tile_view(self):
        """Update tile view"""
        # Alte Widgets loeschen
        for widget in self.tile_scrollable.winfo_children():
            widget.destroy()
        
        # Neue Kacheln erstellen
        for idx, (front_path, back_path) in enumerate(self.images):
            pair_frame = ttk.Frame(self.tile_scrollable, relief=tk.RAISED, borderwidth=2)
            pair_frame.pack(fill=tk.X, padx=5, pady=5)
            pair_frame.pair_index = idx  # Index speichern
            
            # Drag Handle (sichtbarer Bereich oben mit Hinweis)
            drag_handle_frame = tk.Frame(pair_frame, height=20, bg="lightblue", cursor="hand2")
            drag_handle_frame.pack(fill=tk.X)
            drag_handle_frame.pair_index = idx
            drag_label = tk.Label(drag_handle_frame, text="☰ Drag to reorder", 
                                 bg="lightblue", fg="darkblue", font=("Arial", 8))
            drag_label.pack()
            drag_label.pair_index = idx
            
            # Hauptinhalt
            content_frame = ttk.Frame(pair_frame)
            content_frame.pack(fill=tk.BOTH, expand=True)
            
            title_frame = ttk.Frame(content_frame)
            title_frame.pack()
            ttk.Label(title_frame, text=f"Pair {idx + 1}", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)
            # Tausch Button
            swap_btn = ttk.Button(title_frame, text="↔ Swap", width=10,
                               command=lambda i=idx: self.swap_pair_images(i))
            swap_btn.pack(side=tk.LEFT, padx=5)
            
            # Vorderseite Kachel
            front_frame = ttk.Frame(content_frame)
            front_frame.pack(side=tk.LEFT, padx=5, pady=5)
            ttk.Label(front_frame, text="Front").pack()
            front_tile_frame = tk.Frame(front_frame, bg="white", relief=tk.SUNKEN, borderwidth=1)
            front_tile_frame.pack()
            front_tile = tk.Label(front_tile_frame, bg="white")
            front_tile.pack(padx=2, pady=2)
            if front_path:
                self.show_preview(front_path, front_tile, max_size=(120, 120), 
                                pair_index=idx, side='front')
                front_tile.bind("<Button-3>", lambda e, i=idx: self.show_image_menu(e, i, 'front'))
            
            # Rueckseite Kachel
            back_frame = ttk.Frame(content_frame)
            back_frame.pack(side=tk.LEFT, padx=5, pady=5)
            ttk.Label(back_frame, text="Back").pack()
            back_tile_frame = tk.Frame(back_frame, bg="white", relief=tk.SUNKEN, borderwidth=1)
            back_tile_frame.pack()
            back_tile = tk.Label(back_tile_frame, bg="white")
            back_tile.pack(padx=2, pady=2)
            if back_path:
                self.show_preview(back_path, back_tile, max_size=(120, 120),
                                pair_index=idx, side='back')
                back_tile.bind("<Button-3>", lambda e, i=idx: self.show_image_menu(e, i, 'back'))
            else:
                ttk.Label(back_tile_frame, text="(empty)", bg="white").pack(padx=2, pady=2)
            
            # Drag & Drop Events
            def make_drag_start(i):
                return lambda e: self.on_drag_start(e, i)
            def make_drag_motion(i):
                return lambda e: self.on_drag_motion(e, i)
            def make_drag_end(i):
                return lambda e: self.on_drag_end(e, i)
            def make_click_handler(i):
                return lambda e: self.select_pair(i)
            
            # Drag Handle Events
            drag_handle_frame.bind("<Button-1>", make_drag_start(idx))
            drag_handle_frame.bind("<B1-Motion>", make_drag_motion(idx))
            drag_handle_frame.bind("<ButtonRelease-1>", make_drag_end(idx))
            drag_label.bind("<Button-1>", make_drag_start(idx))
            drag_label.bind("<B1-Motion>", make_drag_motion(idx))
            drag_label.bind("<ButtonRelease-1>", make_drag_end(idx))
            
            # Klick Handler fuer Auswahl (nur wenn nicht gedraggt wurde)
            # Nur auf content_frame, nicht auf drag_handle
            content_frame.bind("<Button-1>", make_click_handler(idx))
            for child in content_frame.winfo_children():
                if isinstance(child, (ttk.Frame, tk.Frame, tk.Label, ttk.Label)):
                    child.bind("<Button-1>", make_click_handler(idx))
            
            # Verhindere Auswahl beim Klick auf Drag Handle
            def prevent_selection(e):
                return "break"
            drag_handle_frame.bind("<Button-1>", lambda e: None)  # Wird von drag_start uebernommen
            drag_label.bind("<Button-1>", lambda e: None)  # Wird von drag_start uebernommen
            
            # Rechtsklick-Menue fuer Paar (Loeschen)
            def make_pair_menu_handler(i):
                return lambda e: self.show_pair_menu(e, i)
            pair_frame.bind("<Button-3>", make_pair_menu_handler(idx))
            content_frame.bind("<Button-3>", make_pair_menu_handler(idx))
            
            # Bild-Drag & Drop fuer Tausch zwischen Paaren
            if front_path:
                front_tile.bind("<Button-1>", lambda e, i=idx, s='front': self.on_image_drag_start(e, i, s))
                front_tile.bind("<B1-Motion>", lambda e, i=idx, s='front': self.on_image_drag_motion(e, i, s))
                front_tile.bind("<ButtonRelease-1>", lambda e, i=idx, s='front': self.on_image_drag_end(e, i, s))
                front_tile_frame.bind("<Button-1>", lambda e, i=idx, s='front': self.on_image_drag_start(e, i, s))
                front_tile_frame.bind("<B1-Motion>", lambda e, i=idx, s='front': self.on_image_drag_motion(e, i, s))
                front_tile_frame.bind("<ButtonRelease-1>", lambda e, i=idx, s='front': self.on_image_drag_end(e, i, s))
            
            if back_path:
                back_tile.bind("<Button-1>", lambda e, i=idx, s='back': self.on_image_drag_start(e, i, s))
                back_tile.bind("<B1-Motion>", lambda e, i=idx, s='back': self.on_image_drag_motion(e, i, s))
                back_tile.bind("<ButtonRelease-1>", lambda e, i=idx, s='back': self.on_image_drag_end(e, i, s))
                back_tile_frame.bind("<Button-1>", lambda e, i=idx, s='back': self.on_image_drag_start(e, i, s))
                back_tile_frame.bind("<B1-Motion>", lambda e, i=idx, s='back': self.on_image_drag_motion(e, i, s))
                back_tile_frame.bind("<ButtonRelease-1>", lambda e, i=idx, s='back': self.on_image_drag_end(e, i, s))
    
    def select_pair(self, index):
        """Select pair in preview"""
        if not self.dragging:  # Nur auswaehlen wenn nicht gedraggt wurde
            self.current_pair_index = index
            self.update_previews()
        self.log_debug(f"Pair {index + 1} selected")
    
    def on_drag_start(self, event, index):
        """Drag starten"""
        self.drag_start_index = index
        self.drag_start_y = event.y_root
        self.dragging = False
        self.log_debug(f"Drag started at pair {index + 1}")
    
    def on_drag_motion(self, event, index):
        """Waehrend des Drags"""
        if self.drag_start_index is None:
            return
        
        # Pruefe ob genug bewegt wurde (Drag Threshold)
        if abs(event.y_root - self.drag_start_y) > self.drag_threshold:
            self.dragging = True
            
            # Finde Zielposition
            target_y = event.y_root
            target_index = self.find_drop_position(target_y)
            
            # Visuelles Feedback
            if target_index != self.drag_start_index:
                # Aktualisiere visuelles Feedback
                self.update_drag_feedback(self.drag_start_index, target_index)
    
    def on_drag_end(self, event, index):
        """Drag beenden"""
        if self.drag_start_index is None:
            return
        
        if self.dragging:
            # Finde Zielposition
            target_y = event.y_root
            target_index = self.find_drop_position(target_y)
            
            # Reihenfolge aendern
            if target_index != self.drag_start_index and target_index is not None:
                self.reorder_pairs(self.drag_start_index, target_index)
                self.log_debug(f"Pair {self.drag_start_index + 1} moved to position {target_index + 1}")
        
        # Reset
        self.drag_start_index = None
        self.drag_start_y = None
        self.dragging = False
        
        # Visuelles Feedback zuruecksetzen
        self.update_tile_view()
        self.update_previews()
    
    def find_drop_position(self, y_root):
        """Finde Zielposition basierend auf Y-Koordinate"""
        # Hole alle Pair Frames mit ihren Positionen
        pair_frames = []
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                try:
                    # Hole die Y-Position des Widgets relativ zum Bildschirm
                    widget_y = widget.winfo_rooty()
                    widget_height = widget.winfo_height()
                    pair_frames.append({
                        'index': widget.pair_index,
                        'y': widget_y,
                        'height': widget_height,
                        'center': widget_y + widget_height / 2,
                        'top': widget_y,
                        'bottom': widget_y + widget_height
                    })
                except:
                    # Fallback: verwende Index als Position
                    pair_frames.append({
                        'index': widget.pair_index,
                        'y': widget.pair_index * 200,  # Geschaetzte Hoehe
                        'height': 200,
                        'center': widget.pair_index * 200 + 100,
                        'top': widget.pair_index * 200,
                        'bottom': (widget.pair_index + 1) * 200
                    })
        
        if not pair_frames:
            return self.drag_start_index
        
        # Sortiere nach Y-Position
        pair_frames.sort(key=lambda x: x['y'])
        
        # Finde passende Position basierend auf Mausposition
        for frame_info in pair_frames:
            if frame_info['top'] <= y_root <= frame_info['bottom']:
                # Wenn ueber der Mitte, davor einfügen, sonst danach
                if y_root < frame_info['center']:
                    # Finde Index vor diesem Frame
                    current_pos = pair_frames.index(frame_info)
                    if current_pos > 0:
                        return pair_frames[current_pos - 1]['index']
                    return frame_info['index']
                else:
                    return frame_info['index']
        
        # Wenn ueber allen, zurueck zum ersten
        if y_root < pair_frames[0]['top']:
            return pair_frames[0]['index']
        
        # Wenn unter allen, zurueck zum letzten
        return pair_frames[-1]['index']
    
    def reorder_pairs(self, from_index, to_index):
        """Reorder pairs"""
        if from_index == to_index:
            return
        
        # Paar verschieben
        pair = self.images.pop(from_index)
        self.images.insert(to_index, pair)
        
        # Aktueller Index anpassen
        if self.current_pair_index == from_index:
            self.current_pair_index = to_index
        elif from_index < self.current_pair_index <= to_index:
            self.current_pair_index -= 1
        elif to_index <= self.current_pair_index < from_index:
            self.current_pair_index += 1
    
    def update_drag_feedback(self, from_index, to_index):
        """Visuelles Feedback waehrend des Drags"""
        # Aktualisiere die Kachelansicht mit visuellen Markierungen
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                if widget.pair_index == from_index:
                    widget.config(relief=tk.SUNKEN, borderwidth=3)
                elif widget.pair_index == to_index:
                    widget.config(relief=tk.RIDGE, borderwidth=3)
                else:
                    widget.config(relief=tk.RAISED, borderwidth=2)
    
    def prev_pair(self):
        """Previous pair"""
        if self.images and self.current_pair_index > 0:
            self.current_pair_index -= 1
            self.update_previews()
    
    def next_pair(self):
        """Next pair"""
        if self.images and self.current_pair_index < len(self.images) - 1:
            self.current_pair_index += 1
            self.update_previews()
    
    def show_image_menu(self, event, pair_index, side):
        """Show context menu for image mirroring"""
        menu = tk.Menu(self.root, tearoff=0)
        
        current_mirror = self.image_mirrors.get((pair_index, side), 'none')
        
        # Aktueller Status anzeigen
        status_text = "Current: "
        if current_mirror == 'none':
            status_text += "No mirroring"
        elif current_mirror == 'h':
            status_text += "Horizontal"
        elif current_mirror == 'v':
            status_text += "Vertical"
        elif current_mirror == 'both':
            status_text += "Both"
        menu.add_command(label=status_text, state=tk.DISABLED)
        menu.add_separator()
        
        # Spiegelungsoptionen
        menu.add_command(label="No mirroring", 
                        command=lambda: self.set_image_mirror(pair_index, side, 'none'))
        menu.add_command(label="Mirror horizontally", 
                        command=lambda: self.set_image_mirror(pair_index, side, 'h'))
        menu.add_command(label="Mirror vertically", 
                        command=lambda: self.set_image_mirror(pair_index, side, 'v'))
        menu.add_command(label="Mirror both", 
                        command=lambda: self.set_image_mirror(pair_index, side, 'both'))
        
        # Menue anzeigen
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
    
    def set_image_mirror(self, pair_index, side, mirror_type):
        """Set mirroring for an image"""
        if mirror_type == 'none':
            # Entferne Eintrag wenn keine Spiegelung
            if (pair_index, side) in self.image_mirrors:
                del self.image_mirrors[(pair_index, side)]
        else:
            self.image_mirrors[(pair_index, side)] = mirror_type
        
        side_name = "Front" if side == 'front' else "Back"
        mirror_name = {'none': 'None', 'h': 'Horizontal', 'v': 'Vertical', 'both': 'Both'}[mirror_type]
        self.log_debug(f"Pair {pair_index + 1} {side_name}: mirroring set to '{mirror_name}'")
        
        # Aktualisiere Vorschau
        self.update_previews()
        self.update_tile_view()
    
    def swap_pair_images(self, pair_index):
        """Swap front and back within a pair"""
        if 0 <= pair_index < len(self.images):
            front_path, back_path = self.images[pair_index]
            
            # Tausche die Pfade
            self.images[pair_index] = (back_path, front_path)
            
            # Tausche auch die Spiegelungseinstellungen
            front_mirror = self.image_mirrors.pop((pair_index, 'front'), None)
            back_mirror = self.image_mirrors.pop((pair_index, 'back'), None)
            
            if front_mirror:
                self.image_mirrors[(pair_index, 'back')] = front_mirror
            if back_mirror:
                self.image_mirrors[(pair_index, 'front')] = back_mirror
            
            self.log_debug(f"Pair {pair_index + 1}: front/back swapped")
            
            # Aktualisiere Anzeige
            if self.current_pair_index == pair_index:
                self.update_previews()
            self.update_tile_view()
    
    def show_pair_menu(self, event, pair_index):
        """Show context menu for pair"""
        menu = tk.Menu(self.root, tearoff=0)
        
        menu.add_command(label=f"Delete pair {pair_index + 1}", 
                        command=lambda: self.delete_pair(pair_index))
        menu.add_separator()
        menu.add_command(label="Cancel")
        
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
    
    def delete_pair(self, pair_index):
        """Delete pair"""
        if 0 <= pair_index < len(self.images):
            # Loesche Spiegelungseinstellungen
            for key in list(self.image_mirrors.keys()):
                if key[0] == pair_index:
                    del self.image_mirrors[key]
            
            # Loesche Paar
            self.images.pop(pair_index)
            
            # Aktualisiere Indizes in image_mirrors
            new_mirrors = {}
            for (idx, side), mirror_type in self.image_mirrors.items():
                if idx > pair_index:
                    new_mirrors[(idx - 1, side)] = mirror_type
                elif idx < pair_index:
                    new_mirrors[(idx, side)] = mirror_type
            self.image_mirrors = new_mirrors
            
            # Aktualisiere current_pair_index
            if self.current_pair_index >= len(self.images):
                self.current_pair_index = max(0, len(self.images) - 1)
            elif self.current_pair_index > pair_index:
                self.current_pair_index -= 1
            
            self.log_debug(f"Pair {pair_index + 1} deleted")
            
            # Aktualisiere Anzeige
            self.update_previews()
            self.update_tile_view()
    
    def on_image_drag_start(self, event, pair_index, side):
        """Start image drag"""
        self.image_drag_source = (pair_index, side)
        self.image_drag_start_x = event.x_root
        self.image_drag_start_y = event.y_root
        self.image_dragging = False
        self.log_debug(f"Image drag started: pair {pair_index + 1}, {side}")
    
    def on_image_drag_motion(self, event, pair_index, side):
        """During image drag"""
        if self.image_drag_source is None:
            return
        
        # Pruefe ob genug bewegt wurde
        if hasattr(self, 'image_drag_start_x') and hasattr(self, 'image_drag_start_y'):
            if abs(event.x_root - self.image_drag_start_x) > self.drag_threshold or \
               abs(event.y_root - self.image_drag_start_y) > self.drag_threshold:
                self.image_dragging = True
    
    def on_image_drag_end(self, event, pair_index, side):
        """Finish image drag"""
        if self.image_drag_source is None:
            return
        
        source_pair, source_side = self.image_drag_source
        
        # Finde Ziel-Bild durch Mausposition
        target_pair, target_side = self.find_image_at_position(event.x_root, event.y_root)
        
        # Tausche Bilder wenn Ziel gefunden
        if target_pair is not None and target_side is not None:
            if (source_pair, source_side) != (target_pair, target_side):
                self.swap_images_between_pairs(source_pair, source_side, target_pair, target_side)
        
        # Reset
        self.image_drag_source = None
        self.image_dragging = False
    
    def find_image_at_position(self, x_root, y_root):
        """Find image at mouse position"""
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                widget_y = widget.winfo_rooty()
                widget_height = widget.winfo_height()
                widget_x = widget.winfo_rootx()
                widget_width = widget.winfo_width()
                
                if widget_x <= x_root <= widget_x + widget_width and widget_y <= y_root <= widget_y + widget_height:
                    # Finde ob front oder back - linke Haelfte = front, rechte Haelfte = back
                    pair_idx = widget.pair_index
                    relative_x = x_root - widget_x
                    if relative_x < widget_width / 2:
                        return (pair_idx, 'front')
                    else:
                        return (pair_idx, 'back')
        return (None, None)
    
    def swap_images_between_pairs(self, source_pair, source_side, target_pair, target_side):
        """Swap images between pairs"""
        if (source_pair == target_pair and source_side == target_side) or \
           source_pair >= len(self.images) or target_pair >= len(self.images):
            return
        
        source_front, source_back = self.images[source_pair]
        target_front, target_back = self.images[target_pair]
        
        # Hole Quell- und Ziel-Bilder
        source_image = source_front if source_side == 'front' else source_back
        target_image = target_front if target_side == 'front' else target_back
        
        # Tausche die Bilder
        if source_side == 'front':
            if target_side == 'front':
                self.images[source_pair] = (target_image, source_back)
                self.images[target_pair] = (source_image, target_back)
            else:  # target_side == 'back'
                self.images[source_pair] = (target_image, source_back)
                self.images[target_pair] = (target_front, source_image)
        else:  # source_side == 'back'
            if target_side == 'front':
                self.images[source_pair] = (source_front, target_image)
                self.images[target_pair] = (source_image, target_back)
            else:  # target_side == 'back'
                self.images[source_pair] = (source_front, target_image)
                self.images[target_pair] = (target_front, source_image)
        
        # Tausche auch Spiegelungseinstellungen
        source_mirror = self.image_mirrors.pop((source_pair, source_side), None)
        target_mirror = self.image_mirrors.pop((target_pair, target_side), None)
        
        if source_mirror:
            self.image_mirrors[(target_pair, target_side)] = source_mirror
        if target_mirror:
            self.image_mirrors[(source_pair, source_side)] = target_mirror
        
        self.log_debug(f"Images swapped: pair {source_pair + 1} {source_side} <-> pair {target_pair + 1} {target_side}")
        
        # Aktualisiere Anzeige
        if self.current_pair_index in [source_pair, target_pair]:
            self.update_previews()
        self.update_tile_view()
    
    
    def prepare_image_for_print(self, image_path, mirror=False, pair_index=None, side=None):
        """Prepare image for printing"""
        if not image_path:
            return None
        
        try:
            img = Image.open(image_path)
            
            # Individuelle Spiegelung verwenden wenn vorhanden
            if pair_index is not None and side is not None:
                mirror_type = self.image_mirrors.get((pair_index, side), 'none')
                img = self.apply_mirror(img, mirror_type)
                if mirror_type != 'none':
                    self.log_debug(f"Image mirrored ({mirror_type}): {image_path}")
            # Fallback: Globale Spiegelung (fuer Rueckseite)
            elif mirror:
                img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
                self.log_debug(f"Image mirrored (global): {image_path}")
            
            # Skalierung auf 29.7 cm Breite (A4)
            # Hinweis: Die tatsaechliche Groesse im PDF wird in create_pdf() gesetzt
            # Hier skalieren wir nur fuer Word-Export und Vorschau
            # Fuer PDF wird die Groesse direkt in cm gesetzt
            if self.scale_to_width.get():
                # Berechne Zielgroesse in cm
                margin_cm = self.margin.get()
                available_width_cm = self.target_width - (2 * margin_cm)
                
                # Aspect Ratio beibehalten
                aspect_ratio = img.height / img.width
                target_width_cm = available_width_cm
                target_height_cm = target_width_cm * aspect_ratio
                
                # Pruefen ob Hoehe passt (A4 Hoehe: 21.0 cm)
                max_height_cm = 21.0 - (2 * margin_cm)
                if target_height_cm > max_height_cm:
                    target_height_cm = max_height_cm
                    target_width_cm = target_height_cm / aspect_ratio
                
                # Skaliere Bild fuer bessere Qualitaet (hohe Aufloesung)
                # Verwende 300 DPI als Ziel-Aufloesung
                target_dpi = 300
                target_width_px = int(target_width_cm / 2.54 * target_dpi)
                target_height_px = int(target_height_cm / 2.54 * target_dpi)
                
                img = img.resize((target_width_px, target_height_px), Image.Resampling.LANCZOS)
                self.log_debug(f"Bild skaliert auf {target_width_cm:.2f} x {target_height_cm:.2f} cm ({target_width_px}x{target_height_px} px)")
            
            return img
        except Exception as e:
            self.log_debug(f"Failed to prepare {image_path}: {e}")
            return None
    
    def print_images(self):
        """Print images"""
        if not self.images:
            messagebox.showwarning("Warning", "No images to print.")
            return
        
        try:
            # Temporaeres PDF erstellen
            import tempfile
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_pdf.close()
            
            self.create_pdf(temp_pdf.name)
            
            # PDF oeffnen mit Druckdialog (Windows)
            if sys.platform == 'win32':
                # Oeffne PDF im Standard-Viewer, der dann den Druckdialog zeigt
                os.startfile(temp_pdf.name)
                self.log_debug("PDF opened - use print dialog")
            else:
                # Linux/Mac - versuche mit Druckdialog
                try:
                    # Versuche mit lpr (Linux) oder lp (Mac)
                    if sys.platform == 'darwin':  # Mac
                        os.system(f'open -a "Preview" "{temp_pdf.name}"')
                    else:  # Linux
                        os.system(f'xdg-open "{temp_pdf.name}"')
                except:
                    os.system(f'xdg-open {temp_pdf.name}')
            
        except Exception as e:
            self.log_debug(f"Print failed: {e}")
            messagebox.showerror("Error", f"Print failed: {e}")
    
    def create_pdf(self, filename):
        """Create PDF"""
        # PDF Seitenformat (Standard: Hochformat)
        page_size = landscape(A4) if self.pdf_landscape.get() else A4
        c = canvas.Canvas(filename, pagesize=page_size)
        width, height = page_size
        page_width_cm = 29.7 if self.pdf_landscape.get() else 21.0
        page_height_cm = 21.0 if self.pdf_landscape.get() else 29.7
        margin = self.margin.get() * cm
        
        self.log_debug(f"Creating PDF: {filename}")
        
        for idx, (front_path, back_path) in enumerate(self.images):
            # Vorderseite
            # Lade Originalbild fuer Groessenberechnung
            if front_path:
                base_img = self.load_base_image(front_path, pair_index=idx, side='front',
                                               mirror=False, trim=self.auto_trim.get())
                if self.scale_to_width.get():
                    margin_cm = self.margin.get()
                    available_width_cm = page_width_cm - (2 * margin_cm)
                    available_height_cm = page_height_cm - (2 * margin_cm)
                    img_width_cm, img_height_cm = self.compute_target_size_cm(
                        base_img, available_width_cm, available_height_cm
                    )
                    export_img = self.prepare_export_image(base_img, img_width_cm, img_height_cm)
                    img_width = img_width_cm * cm
                    img_height = img_height_cm * cm
                else:
                    export_img = base_img
                    dpi = 72
                    img_width = (export_img.width / dpi) * 2.54 * cm
                    img_height = (export_img.height / dpi) * 2.54 * cm
                
                # Zentrieren
                x = (width - img_width) / 2
                y = (height - img_height) / 2
                
                c.drawImage(ImageReader(export_img), x, y, width=img_width, height=img_height)
                self.log_debug(f"Front {idx + 1} added: {img_width/cm:.2f} x {img_height/cm:.2f} cm")
            
            c.showPage()
            
            # Rueckseite
            if back_path:
                use_global_mirror = self.mirror_back.get() if (idx, 'back') not in self.image_mirrors else False
                base_back_img = self.load_base_image(back_path, pair_index=idx, side='back',
                                                    mirror=use_global_mirror, trim=self.auto_trim.get())
                if self.scale_to_width.get():
                    margin_cm = self.margin.get()
                    available_width_cm = page_width_cm - (2 * margin_cm)
                    available_height_cm = page_height_cm - (2 * margin_cm)
                    img_width_cm, img_height_cm = self.compute_target_size_cm(
                        base_back_img, available_width_cm, available_height_cm
                    )
                    export_back_img = self.prepare_export_image(base_back_img, img_width_cm, img_height_cm)
                    img_width = img_width_cm * cm
                    img_height = img_height_cm * cm
                else:
                    export_back_img = base_back_img
                    dpi = 72
                    img_width = (export_back_img.width / dpi) * 2.54 * cm
                    img_height = (export_back_img.height / dpi) * 2.54 * cm
                
                # Zentrieren
                x = (width - img_width) / 2
                y = (height - img_height) / 2
                
                c.drawImage(ImageReader(export_back_img), x, y, width=img_width, height=img_height)
                self.log_debug(f"Back {idx + 1} added: {img_width/cm:.2f} x {img_height/cm:.2f} cm")
            else:
                self.log_debug(f"Back {idx + 1} is empty")
            
            c.showPage()
        
        c.save()
        self.log_debug(f"PDF saved: {filename}")
    
    def open_file(self, filepath):
        """Open file with default app (cross-platform)"""
        try:
            if sys.platform == 'win32':
                try:
                    os.startfile(filepath)
                except Exception:
                    # Fallback: cmd start
                    subprocess.Popen(["cmd", "/c", "start", "", filepath], shell=True)
            elif sys.platform == 'darwin':  # Mac
                os.system(f'open "{filepath}"')
            else:  # Linux
                os.system(f'xdg-open "{filepath}"')
            self.log_debug(f"File opened: {filepath}")
            return True
        except Exception as e:
            self.log_debug(f"Failed to open file: {e}")
            return False
    
    def export_pdf(self):
        """Save as PDF"""
        if not self.images:
            messagebox.showwarning("Warning", "No images to export.")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                self.create_pdf(filename)
                
                # Automatisch oeffnen wenn aktiviert
                if self.auto_open_export.get():
                    self.open_file(filename)
            except Exception as e:
                self.log_debug(f"PDF export failed: {e}")
                messagebox.showerror("Error", f"Save failed: {e}")
    
    def export_word(self):
        """Save as Word document"""
        if not self.images:
            messagebox.showwarning("Warning", "No images to export.")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                doc = Document()
                section = doc.sections[0]
                # A4 im Querformat, damit 29.7 cm die Seitenbreite ist
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
                # Margen explizit setzen, damit die Breite korrekt skaliert wird
                section.left_margin = Cm(self.margin.get())
                section.right_margin = Cm(self.margin.get())
                section.top_margin = Cm(self.margin.get())
                section.bottom_margin = Cm(self.margin.get())
                available_width_cm = (section.page_width - section.left_margin - section.right_margin) / Cm(1)
                available_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / Cm(1)
                
                self.log_debug(f"Creating Word document: {filename}")
                
                for idx, (front_path, back_path) in enumerate(self.images):
                    # Vorderseite auf eigener Seite
                    if front_path:
                        base_front_img = self.load_base_image(front_path, pair_index=idx, side='front',
                                                             mirror=False, trim=self.auto_trim.get())
                        front_img = base_front_img
                        if front_img:
                            # Bild in Bytes speichern
                            img_bytes = io.BytesIO()
                            front_img.save(img_bytes, format='PNG')
                            img_bytes.seek(0)
                            
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            
                            # Berechne korrekte Groesse wie im PDF
                            if self.scale_to_width.get():
                                img_width_cm, img_height_cm = self.compute_target_size_cm(
                                    front_img, available_width_cm, available_height_cm
                                )
                                # Setze Breite und Hoehe
                                run.add_picture(img_bytes, width=Cm(img_width_cm), height=Cm(img_height_cm))
                                self.log_debug(f"Front {idx + 1} added: {img_width_cm:.2f} x {img_height_cm:.2f} cm")
                            else:
                                # Originalgroesse verwenden (auf Seitenbreite)
                                run.add_picture(img_bytes, width=Cm(available_width_cm))
                                self.log_debug(f"Front {idx + 1} added (page width)")
                    
                    # Seitenumbruch fuer Rueckseite (beidseitiger Druck)
                    doc.add_page_break()
                    
                    # Rueckseite auf eigener Seite
                    if back_path:
                        # Verwende individuelle Spiegelung oder globale Einstellung
                        use_global_mirror = self.mirror_back.get() if (idx, 'back') not in self.image_mirrors else False
                        base_back_img = self.load_base_image(back_path, pair_index=idx, side='back',
                                                            mirror=use_global_mirror, trim=self.auto_trim.get())
                        back_img = base_back_img
                        if back_img:
                            img_bytes = io.BytesIO()
                            back_img.save(img_bytes, format='PNG')
                            img_bytes.seek(0)
                            
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            
                            # Berechne korrekte Groesse wie im PDF
                            if self.scale_to_width.get():
                                img_width_cm, img_height_cm = self.compute_target_size_cm(
                                    back_img, available_width_cm, available_height_cm
                                )
                                # Setze Breite und Hoehe
                                run.add_picture(img_bytes, width=Cm(img_width_cm), height=Cm(img_height_cm))
                                self.log_debug(f"Back {idx + 1} added: {img_width_cm:.2f} x {img_height_cm:.2f} cm")
                            else:
                                # Originalgroesse verwenden (auf Seitenbreite)
                                run.add_picture(img_bytes, width=Cm(available_width_cm))
                                self.log_debug(f"Back {idx + 1} added (page width)")
                    else:
                        # Leere Seite wenn keine Rueckseite
                        paragraph = doc.add_paragraph()
                        paragraph.add_run("(No back side)")
                        self.log_debug(f"Back {idx + 1} is empty")
                    
                    # Seitenumbruch nur wenn nicht das letzte Paar
                    if idx < len(self.images) - 1:
                        doc.add_page_break()
                
                doc.save(filename)
                self.log_debug(f"Word document saved: {filename}")
                
                # Automatisch oeffnen wenn aktiviert
                if self.auto_open_export.get():
                    self.open_file(filename)
            except Exception as e:
                self.log_debug(f"Word export failed: {e}")
                messagebox.showerror("Error", f"Save failed: {e}")


def main():
    if DND_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    
    app = DruckManager(root)
    root.mainloop()


if __name__ == "__main__":
    main()
